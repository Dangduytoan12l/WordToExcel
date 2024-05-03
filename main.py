import os
import re
import docx
import win32com.client as win32
from pypandoc import convert_file
from utils import CFL, create_quiz, extract_format_text, split_options, is_option, is_question, process_formats

def format_file(file_path: str, del_list: list, selected_options: list) -> list:
    """
    Format a document file to DOCX, extract formatted text, and return relevant information.
    
    Args:
        file_path (str): The path to the document file.
        del_list (list): A list to store the paths of temporary files.
        selected_options (list): A list of selected options for formatting.
        
    Returns:
        list: Containing the path to the converted DOCX file, a list of highlighted text, and a list of paths of temporary files.
    """

    def convert_to_docx(docx_path: str, name: str, del_list: list) -> list:
        global temp_path_docx
        temp_name = f'wteTemp{name}'
        
        # Load the DOCX document using pypandoc
        convert_file(docx_path, 'plain', extra_args=['--wrap=none'], outputfile=f'{temp_name}.txt')
        document = docx.Document()
        
        # Read the text from the file and replace soft returns with paragraph marks
        with open(f'{temp_name}.txt', 'r', encoding='utf-8') as file:
            text = file.readlines()
        
        for line in text:
            document.add_paragraph(line)
        
        document.save(f'{temp_name}.docx')
        temp_path_docx = os.path.abspath(f'{temp_name}.docx')
        del_list.append(temp_path_docx)
        return del_list

    # Function to extract formatted text
    def extract_original_format(file_path: str, selected_options: list) -> list:
        highlights = []
        document = docx.Document(file_path)
        # Append the highlighted text
        for paragraph in document.paragraphs:
            highlighted_text = extract_format_text(paragraph, selected_options)
            if highlighted_text is not None:
                match = re.match(r'^[a-dA-D]', highlighted_text)
                if "A,B,C,D" in selected_options and match:
                    highlights.append(highlighted_text)
                # Regex to extract the correct answer with no white space
                else:
                    highlights.append(CFL(re.sub(f'{match}. ', '', highlighted_text)))
        return highlights

    # Split the file path into name and extension
    name, ext = os.path.splitext(os.path.basename(file_path))
    abs_file_path = os.path.abspath(file_path)
    
    if ext == ".doc":
        # Convert .doc to .docx
        temp_name = f"wteDocTemp{name}"
        temp_path = os.path.abspath(f"{temp_name}.docx")
        
        word = win32.gencache.EnsureDispatch('Word.Application')
        doc = word.Documents.Open(abs_file_path)
        doc.Activate()
        
        word.ActiveDocument.SaveAs(temp_name, FileFormat=win32.constants.wdFormatXMLDocument)
        doc.Close(False)
        word.Quit()
        
        #Delete the temporary .docx file
        del_list = convert_to_docx(temp_path, name, del_list)
        del_list.append(temp_path)
        highlights = extract_original_format(temp_path, selected_options)
        
        return temp_path_docx, highlights, del_list

    elif ext == ".docx":
        del_list = convert_to_docx(abs_file_path, name, del_list)
        highlights = extract_original_format(file_path, selected_options)
        return temp_path_docx, highlights, del_list   
    
    return None, None, None 

# Function to process questions and options
def question_create(doc, current_question: str, current_options: list, highlights: list, data: list, platform: str, selected_options: list, question_numbers: int) -> int:
    """
    Process a document to create quiz questions and options based on specific formatting.
    The document structure and formatting rules must align with the processing logic for accurate results.
    
    Parameters:
    - doc: The document object to process.
    - current_question: The current question being processed.
    - current_options: The list of current options being processed.
    - highlights: The list of highlights being processed.
    - data: The list of data being processed.
    - platform: The platform for which the quiz is being created.
    - selected_options: The list of selected options being processed.
    - question_numbers: The current question number.
    
    Returns:
    - int: The updated question number.
    """

    def last_question(current_question: str, current_options: list, highlights: list, data: list, platform: str, selected_options: list, question_numbers: int) -> None:
        if current_question and len(current_options) > 0:
            current_question, current_options = process_formats(current_question, current_options, selected_options, question_numbers)
            create_quiz(data, current_question, current_options, highlights, platform, selected_options)
    
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        #The second condition is to handle multiple line questions, index-2 is because of how the current question is updated
        if is_question(text) and not is_question(doc.paragraphs[index - 2].text.strip()):
            if current_question and len(current_options) > 0:
                current_question, current_options = process_formats(current_question, current_options, selected_options, question_numbers)
                question_numbers += 1
                create_quiz(data, current_question, current_options, highlights, platform, selected_options)
            current_options.clear()  # Clear the options list for the new questions
            current_question = text
        elif current_question:
            if is_option(text):
                for option in split_options(text):
                    current_options.append(option)
            elif text.strip() and not is_option(doc.paragraphs[index-2].text.strip()): 
                current_question += '\n'+text
    # Process the last question
    question_numbers += 1
    last_question(current_question, current_options, highlights, data, platform, selected_options, question_numbers)
    return question_numbers

